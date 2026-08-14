from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Auto-Tune后续研发路线与实施评估_技术Leader版_20260814.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
NAVY = "0B2545"
GRAY = "5B6573"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4CE"
PALE_RED = "FDE9E7"
GREEN = "E2F0D9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=10.5, bold=False, color="000000", italic=False):
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, bold=False, color="000000", size=10.5, after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(text), size=10.2)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    set_run(p.add_run(text), size=10.2)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_table(doc, headers, rows, widths, font_size=9.0, fills=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], LIGHT)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(h), size=font_size, bold=True, color=NAVY)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if fills and r_idx < len(fills) and fills[r_idx]:
                set_cell_shading(cells[i], fills[r_idx])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_callout(doc, title, body, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(title), size=10.5, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08
    set_run(p2.add_run(body), size=9.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def page_break(doc):
    doc.add_page_break()


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK, 8, 4),
    ):
        s = doc.styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        s = doc.styles[name]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(10.2)


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(hp.add_run("Auto-Tune｜后续研发路线与实施评估｜内部评审稿"), size=8.5, color=GRAY)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(fp.add_run("内部资料  ·  2026-08-14  ·  第 "), size=8.5, color=GRAY)
    add_page_field(fp)
    set_run(fp.add_run(" 页"), size=8.5, color=GRAY)

    # Cover / memo masthead
    add_text(doc, "技术方案评审与资源申请", bold=True, color=BLUE, size=10.5, after=8)
    add_text(doc, "Auto-Tune 视觉模型自动调优平台", bold=True, color=NAVY, size=24, after=4)
    add_text(doc, "后续研发路线与实施评估", bold=True, color=DARK, size=18, after=16)
    add_text(doc, "面向技术 Leader：完整范围、技术难度、资源需求、工期与风险门槛", color=GRAY, size=11.5, after=18)
    meta = [
        ("文档日期", "2026-08-14"),
        ("当前阶段", "一期：LLM 仅在受控范围内调整超参数并触发训练"),
        ("规划范围", "本地可靠性 → 在线多人平台 → 多任务/版本 → 结构实验 → CVAT 集成"),
        ("估算口径", "工程人日；不含外部采购、生产级运维值守和大规模训练资源排队"),
        ("建议评审结论", "批准分阶段推进；每阶段以验收门槛决定是否进入下一阶段"),
    ]
    add_table(doc, ["项目", "内容"], meta, [1800, 7560], 9.4)
    add_callout(doc, "一句话结论", "方向可行，但不能把后续功能当作在现有单机代码上的简单叠加。建议先完成可靠性与领域接口固化，再建设在线队列和数据隔离；模型结构解析与结构自调整应作为独立一期共同交付，并采用注册表、模板和人工审批约束。", PALE_GOLD)
    add_text(doc, "依据：现有方案审查、implementation_plan_20260814.md、roadmap_20260814.md、当前代码修复结果，以及 81 项自动化测试与 1 epoch 真实训练冒烟验证。", color=GRAY, size=8.8, after=0)

    page_break(doc)
    add_heading(doc, "1. 管理摘要", 1)
    add_callout(doc, "建议决策", "按 A0 → A1/A2/A3 → B → C1/C2/C3 的顺序实施。当前一期维持“只调整超参数”的边界；结构解析和结构自调整在后续同一期交付，但必须与超参数 Agent 分离动作空间。", GREEN)
    add_heading(doc, "1.1 关键判断", 2)
    for t in [
        "项目已经具备数据分析、训练诊断、超参数调优闭环和 Web 操作入口，适合继续演进，不建议推倒重写。",
        "Module A/B/C 的理念可复用，但接口、事实语义、审计与持久化必须先固化，不能直接视为稳定平台内核。",
        "在线多人平台是最大工程拐点：需要数据库、权限隔离、独立 Worker、GPU 原子租约、事件持久化及恢复机制。",
        "模型结构能力属于高难度实验平台，而不是普通配置功能；价值高，但应晚于可靠性和执行平台。",
        "CVAT 应通过 API 解耦集成，不应合并代码库或让 CVAT 直接管理训练进程。",
    ]:
        add_bullet(doc, t)
    add_heading(doc, "1.2 总体投入判断", 2)
    add_table(doc, ["口径", "估算", "说明"], [
        ("剩余完整研发工作量", "约 46–82 人日", "按当前 A0 大部完成后计算；含 A1/A2/A3、B、C1/C2/C3"),
        ("单人串行日历周期", "约 10–18 周", "含联调、评审、缓冲；不建议用于在线平台与结构平台并行目标"),
        ("2 人核心团队", "约 7–12 周", "平台工程与算法/训练工程部分并行"),
        ("3 人跨职能团队", "约 5–9 周", "后端/平台、ML、前端/QA 并行；仍受 GPU 与验收依赖约束"),
    ], [2200, 1800, 5360], 9.2)
    add_text(doc, "注：日历周期为规划估算，不是承诺上线日期；外部环境、数据质量、GPU 可用性和 Ultralytics/CVAT 兼容性会影响区间。", color=GRAY, size=8.8)

    page_break(doc)
    add_heading(doc, "2. 当前基线与已验证能力", 1)
    add_heading(doc, "2.1 现有能力", 2)
    add_table(doc, ["领域", "当前能力", "成熟度判断"], [
        ("Module A 数据分析", "图像质量、bbox 几何、类别分布、聚类与报告", "可用；多任务标签仍需适配"),
        ("Module B 训练诊断", "数值摘要、规则诊断、LLM 文本解释、视觉分析", "可用；需持续提升事实一致性"),
        ("Module C 自动调优", "感知→决策→护栏→训练→探针→复盘", "一期闭环已具备；需强化审计和异常策略"),
        ("Web UI", "目录选择、分析、训练、SSE、结果展示、报告", "单机可用；多人和恢复能力不足"),
        ("存储/调度", "JSON 文件、本地 subprocess、单会话 SSE", "仅适合单机原型"),
    ], [1900, 4660, 2800], 8.8)
    add_heading(doc, "2.2 本轮已完成的可靠性修复", 2)
    for t in [
        "建立统一超参数注册表，Prompt、Guardrails、执行器共享参数语义。",
        "统一校验 hyperparameter_changes 与 training_overrides，Guardrails 清洗值成为唯一可执行值。",
        "拦截未知参数、类型错误和 optimizer=auto 与显式学习率/动量冲突。",
        "修复 Module B 嵌套配置、Early Stopping 读取与训练前后指标差值反馈。",
        "按当前 Python 环境解析 YOLO 入口，减少 PATH 偶然性。",
        "训练完成后同步 KPI、结果载荷和关键日志；保留 ZIP 仅作为安全兼容路径，正式操作改为直接选择目录。",
    ]:
        add_bullet(doc, t)
    add_callout(doc, "验证证据", "自动化测试：81 passed，2 warnings。真实训练：train38，epochs=1；训练完成事件成功返回 precision、recall、mAP50、mAP50-95。该结果证明短流程可运行，但不等价于模型质量达标。", GREEN)

    page_break(doc)
    add_heading(doc, "3. 目标架构与边界", 1)
    add_heading(doc, "3.1 四个稳定领域", 2)
    add_table(doc, ["领域", "职责", "关键对象"], [
        ("数据域", "数据集接入、不可变快照、质量分析", "Dataset、Snapshot、QualityReport"),
        ("训练域", "任务、参数、GPU/进程、产物", "TrainingJob、Lease、Checkpoint、Artifact"),
        ("分析域", "数值事实、规则诊断、LLM/视觉解释", "Facts、Issue、AnalysisReport"),
        ("实验域", "超参数/结构实验、对比、最优模型", "Experiment、Variant、Decision、Score"),
    ], [1600, 4400, 3360], 9.0)
    add_heading(doc, "3.2 一期与后续结构一期的边界", 2)
    add_table(doc, ["能力", "一期：超参数", "后续结构一期"], [
        ("LLM 动作", "仅建议白名单超参数", "仅选择注册模块和合法模板变体"),
        ("执行输入", "Guardrails 输出的 sanitized 参数", "结构生成器验证后的结构指纹/YAML"),
        ("自由度", "每轮 1–2 个相关参数", "优先单结构维度或小组合消融"),
        ("审批", "可配置自动执行", "默认人工审批，实验与生产分离"),
        ("禁止事项", "命令、任意参数、绕过护栏", "任意 import、任意 Python、任意 YAML tag"),
    ], [1700, 3830, 3830], 8.8)
    add_callout(doc, "核心安全原则", "Python 产生事实，规则系统进行确定性判断，LLM 负责解释与有限决策。LLM 不直接生成或执行命令，也不成为训练参数或模型结构的最终可信源。", PALE_RED)
    add_heading(doc, "3.3 本地与在线共存", 2)
    add_text(doc, "本地和在线共享领域服务及适配器；持久化、认证、对象存储和调度通过接口替换。本地可使用单用户模式，在线必须开启用户、权限、配额、队列和审计。")

    page_break(doc)
    add_heading(doc, "4. 分阶段研发路线", 1)
    add_table(doc, ["阶段", "目标", "难度", "工作量", "进入条件"], [
        ("A0", "闭环可靠性收尾", "中高", "剩余约 2–4 人日", "当前代码基线"),
        ("A1", "任务/YOLO 版本适配框架", "中", "2–4 人日", "A0 参数语义稳定"),
        ("A2", "Linux 与部署基线", "中", "1–2 人日", "目标服务器可用"),
        ("A3", "本地交互体验", "中", "1–3 人日", "目录边界确定"),
        ("B", "在线多人内部可用版", "高", "15–25 人日", "A0、A2 完成"),
        ("C1", "分割/OBB/分类/v5 等扩展", "中高", "8–15 人日", "A1 完成，业务优先级明确"),
        ("C2", "结构解析 + 受控结构自调整", "很高", "12–20 人日", "实验域、Worker、资源预算可用"),
        ("C3", "CVAT 解耦集成", "高", "8–15 人日", "B 队列/API 稳定"),
    ], [760, 3150, 900, 1450, 3100], 8.3)
    add_heading(doc, "4.1 推荐顺序", 2)
    for t in [
        "先完成 A0：确保任何进入训练的参数都可解释、可校验、可审计。",
        "A1、A2、A3 可有限并行，但 A1 不应早于 A0 参数注册表稳定。",
        "B 先交付内部可用版，再评估生产化稳定性周期。",
        "C1 按真实业务需求逐项交付，避免一次性泛化所有任务。",
        "C2 作为独立结构实验一期；C3 复用 B 的任务队列和模型产物接口。",
    ]:
        add_bullet(doc, t)

    page_break(doc)
    add_heading(doc, "5. 阶段 A：本地闭环与扩展底座", 1)
    add_heading(doc, "5.1 A0 可靠性收尾（剩余约 2–4 人日）", 2)
    add_table(doc, ["工作包", "内容", "验收"], [
        ("审计闭环", "记录原始建议、候选值、护栏值、实际命令、指标差值和错误", "参数审计一致率 100%"),
        ("数据快照", "训练划分不移动原始文件，生成不可变快照或临时目录", "原始数据可验证未变化"),
        ("密钥治理", "仓库历史轮换，环境变量/本地私密配置", "仓库无有效密钥"),
        ("UI/日志", "默认关键指标，完整日志折叠；历史记录统一", "普通训练和调优均可追溯"),
        ("失败策略", "API/Schema/非法参数失败时不启动训练", "负向用例全部阻断"),
    ], [1800, 4760, 2800], 8.8)
    add_heading(doc, "5.2 A1 任务/版本适配框架（2–4 人日）", 2)
    add_bullet(doc, "定义 TaskAdapter：detect、segment、obb、classify；定义 ModelAdapter：YOLOv8、YOLO11、后续经验证版本、YOLOv5。")
    add_bullet(doc, "适配器统一模型目录、命令构造、指标列映射、评分公式和参数注册表。")
    add_bullet(doc, "先保持 v8 Detect 回归，再接入 v11 Detect；任何新版本先验证权重、CLI、指标列和依赖。")
    add_heading(doc, "5.3 A2/A3（2–5 人日，可部分并行）", 2)
    add_bullet(doc, "Linux：环境文件、CUDA/PyTorch 匹配、路径/编码清理、health/readiness、运行环境指纹。")
    add_bullet(doc, "交互：正式流程直接选择受控目录；展示校验进度、失败原因和重新分析；统一训练/分析/调优状态。")

    page_break(doc)
    add_heading(doc, "6. 阶段 B：在线多人平台（15–25 人日）", 1)
    add_callout(doc, "难度判断：高", "这不是简单增加登录页。真正难点是训练进程与 Web 解耦、GPU 资源不重复分配、服务重启后的任务对账、数据权限隔离以及事件可恢复。", PALE_GOLD)
    add_table(doc, ["工作包", "工作量", "主要交付", "主要难点"], [
        ("B1 数据库/仓储", "3–5 人日", "用户、数据集、快照、任务、事件、报告、产物、GPU 租约", "迁移、幂等、索引、备份恢复"),
        ("B2 认证与隔离", "3–5 人日", "登录、角色、owner/role 校验、配额、产物生命周期", "越权、路径穿越、跨用户访问"),
        ("B3 Worker/GPU 调度", "5–8 人日", "状态机、原子租约、取消、超时、重试、恢复", "并发竞争、孤儿进程、Checkpoint"),
        ("B4 SSE/历史", "2–4 人日", "事件游标、重放、日志分页、任务对比", "断线恢复、事件顺序、一致性"),
        ("B5 部署运维", "2–3 人日", "HTTPS、健康检查、监控、结构化日志、容量预警", "环境一致性、故障定位"),
    ], [1600, 1300, 3660, 2800], 8.5)
    add_heading(doc, "6.1 必须实现的任务状态机", 2)
    add_text(doc, "queued → preparing → running → analyzing → completed / failed / cancelled。状态改变必须持久化并产生带递增游标的事件；服务重启后需对账数据库、PID、Checkpoint 与 GPU 租约。")
    add_heading(doc, "6.2 阶段 B 验收门槛", 2)
    for t in [
        "多用户资源隔离测试全部通过，任何 API 均执行 owner/role 校验。",
        "GPU 同时分配冲突数为 0；取消、超时和失败重试不会遗留孤儿进程。",
        "服务重启后任务状态可对账，SSE 可从最后游标恢复。",
        "任一模型产物可追溯到数据快照、参数、代码/依赖版本和实验记录。",
    ]:
        add_bullet(doc, t)

    page_break(doc)
    add_heading(doc, "7. 阶段 C1：任务类型与旧版本扩展（8–15 人日）", 1)
    add_table(doc, ["扩展项", "难度", "建议工作量", "关键改造"], [
        ("YOLO11 Detect", "中", "1–2 人日", "版本适配、权重/指标回归"),
        ("分割", "中高", "2–4 人日", "多边形标签、mask/box 指标、独立评分"),
        ("OBB", "中高", "2–3 人日", "旋转框几何、角度/方向统计、可视化"),
        ("分类", "高", "3–5 人日", "目录布局、top1/top5、评分与诊断重写"),
        ("YOLOv5", "中高", "2–3 人日", "独立命令适配、结果列/loss 映射、依赖边界"),
    ], [1600, 1000, 1500, 5260], 8.8)
    add_heading(doc, "7.1 推荐业务优先级", 2)
    add_text(doc, "建议默认顺序：YOLO11 Detect → 分割 → OBB → 分类 / YOLOv5。最终顺序应由真实数据和客户需求决定；YOLOv5 仅在遗留模型必须继续训练时接入。")
    add_heading(doc, "7.2 设计约束", 2)
    for t in [
        "不同任务必须拥有独立的数据验证、指标映射和评分公式，不能复用 Detect 的错误类型。",
        "每接入一个任务，先构建最小真实数据集回归，再开放 UI 入口。",
        "新 YOLO 版本不是增加下拉选项；必须验证当前 Ultralytics 包、模型名、权重格式和 results.csv。",
    ]:
        add_bullet(doc, t)

    page_break(doc)
    add_heading(doc, "8. 阶段 C2：模型结构解析与受控结构自调整（12–20 人日）", 1)
    add_callout(doc, "产品评价", "两个功能应一起做，但内部必须拆成“可信结构事实”和“受控结构决策”两层。结构解析本身有独立价值；结构自调整只有在解析、注册表、资源预检和消融记录完整时才可信。", GREEN)
    add_table(doc, ["子模块", "工作量", "交付内容", "难度"], [
        ("C2.1 结构解析器", "3–5 人日", "Backbone/Neck/Head、连接、通道、stride、参数量、FLOPs、显存估计、权重匹配率、结构指纹", "高"),
        ("C2.2 模块注册表", "2–4 人日", "测试过的 Attention/Neck/Conv/Loss；输入输出、任务/版本、资源成本", "高"),
        ("C2.3 变体生成器", "3–5 人日", "合法模板、张量/兼容性/资源预检、baseline/单变量/小组合、失败回退", "很高"),
        ("C2.4 结构 Agent", "2–3 人日", "独立动作空间、资源预算、单维度优先、人工审批", "很高"),
        ("C2.5 实验 UI/审计", "2–3 人日", "结构图、变体差异、消融对比、审批和最佳模型保护", "高"),
    ], [1750, 1250, 5260, 1100], 8.2)
    add_heading(doc, "8.1 为什么技术难度高", 2)
    for t in [
        "结构合法不等于可训练：通道、尺度、张量形状、预训练权重兼容和显存都可能失败。",
        "Ultralytics 内部模块与版本耦合，升级可能改变解析方式、注册入口和权重加载行为。",
        "单次指标提升不能证明结构收益，需要 baseline、消融和重复实验控制随机性。",
        "自由生成结构的搜索空间过大且不可审计，必须限定为注册模块和模板动作。",
    ]:
        add_bullet(doc, t)
    add_heading(doc, "8.2 验收门槛", 2)
    add_text(doc, "每个结构变体必须可解释、可复现、可回退；训练前通过兼容性和资源预检；实验记录具备 baseline 与消融证据；结构 Agent 默认不可直接覆盖当前最佳模型。")

    page_break(doc)
    add_heading(doc, "9. 阶段 C3：CVAT 解耦集成（8–15 人日）", 1)
    add_heading(doc, "9.1 推荐边界", 2)
    add_table(doc, ["CVAT 负责", "Auto-Tune 负责", "接口契约"], [
        ("项目、任务、标注、数据版本", "数据分析、训练、调优、模型交付", "快照 ID、类别映射、来源版本"),
        ("人工审核与标注流程", "训练队列、GPU、Checkpoint、指标", "任务状态、日志、指标、产物"),
        ("自动标注能力调用", "注册经审批模型", "服务 Token/OAuth、最小权限"),
    ], [2800, 2800, 3760], 8.8)
    add_heading(doc, "9.2 关键原则", 2)
    for t in [
        "CVAT 不直接控制训练子进程；Auto-Tune 不直接修改 CVAT 标注。",
        "数据通过固化快照交接，保证训练期间输入不漂移。",
        "训练队列复用阶段 B，不重复建设。",
        "模型回传必须经过审批、版本登记和兼容性检查。",
    ]:
        add_bullet(doc, t)
    add_callout(doc, "主要风险", "CVAT 版本、导出格式、类别映射、Token 权限和自动标注部署方式都可能影响联调周期。建议先以单向导出→训练为最小闭环，再增加模型回传。", PALE_GOLD)

    page_break(doc)
    add_heading(doc, "10. 技术难度与风险矩阵", 1)
    add_table(doc, ["风险", "概率", "影响", "应对策略", "决策门槛"], [
        ("GPU 并发重复分配", "中", "很高", "数据库原子租约、心跳、超时回收、冲突测试", "B 上线前必须为 0"),
        ("用户数据越权/路径穿越", "中", "很高", "owner/role 全接口校验、根目录白名单、负向测试", "安全测试全通过"),
        ("训练进程重启失联", "高", "高", "独立 Worker、PID/Checkpoint 对账、幂等恢复", "故障演练通过"),
        ("LLM 输出污染执行", "中", "高", "Schema、白名单、Guardrails 唯一执行值、失败不训练", "非法参数进入命令=0"),
        ("结构变体不可训练", "高", "高", "模板/注册表、张量和显存预检、失败回退", "C2 预检覆盖关键模块"),
        ("版本升级破坏适配", "中", "中高", "依赖锁定、适配器契约、真实数据回归", "升级前后基线一致"),
        ("数据集训练中漂移", "中", "高", "不可变快照、内容指纹、只读输入", "快照审计可复现"),
        ("CVAT 联调复杂度", "中", "中高", "API 解耦、最小闭环、契约测试", "先单向后双向"),
    ], [1600, 700, 800, 4260, 2000], 7.9)
    add_heading(doc, "10.1 需要 Leader 提前确认的决策", 2)
    for t in [
        "在线平台目标是内部可用还是生产级 SLA；两者的运维与稳定性投入不同。",
        "优先支持的业务任务类型及真实验收数据集。",
        "GPU 数量、型号、共享策略和训练预算。",
        "结构实验是否允许自动执行，还是默认必须人工审批。",
        "CVAT 首期只做数据导出，还是同时要求模型回传与自动标注。",
    ]:
        add_bullet(doc, t)

    page_break(doc)
    add_heading(doc, "11. 人力配置与日历排期", 1)
    add_heading(doc, "11.1 推荐角色", 2)
    add_table(doc, ["角色", "主要职责", "建议投入"], [
        ("技术负责人/架构", "阶段门槛、领域接口、风险决策、代码评审", "0.2–0.4 FTE 全程"),
        ("后端/平台工程师", "数据库、认证、Worker、GPU 调度、SSE、部署", "1 FTE，阶段 B 主力"),
        ("ML/训练工程师", "参数注册、任务适配、结构解析/变体、训练验证", "1 FTE，A/C 主力"),
        ("前端工程师", "目录流程、任务/队列/实验 UI、错误反馈", "0.3–0.6 FTE"),
        ("QA/DevOps", "并发、安全、故障恢复、Linux/CUDA 与发布验证", "0.3–0.6 FTE"),
    ], [1900, 4760, 2700], 8.8)
    add_heading(doc, "11.2 团队规模场景", 2)
    add_table(doc, ["团队方案", "预计周期", "适用场景", "主要代价"], [
        ("1 名全栈/ML 工程师", "10–18 周", "继续原型和顺序交付", "上下文切换大，在线平台和结构平台风险集中"),
        ("2 名核心工程师", "7–12 周", "推荐最低配置", "前端、QA、DevOps 仍需兼职支持"),
        ("3 名跨职能核心成员", "5–9 周", "希望较快形成内部可用平台", "需要技术负责人严格管理接口与并行依赖"),
    ], [2100, 1500, 2960, 2800], 8.7)
    add_text(doc, "上述周期以 5 个工作日/周估算，并包含约 15%–25% 的联调和风险缓冲。若要求生产级高可用、安全审计或多节点 GPU 调度，应单独追加稳定化周期。", color=GRAY, size=8.8)

    page_break(doc)
    add_heading(doc, "12. 建议里程碑与评审点", 1)
    add_table(doc, ["里程碑", "建议时间点（2 人团队）", "交付物", "Go/No-Go 条件"], [
        ("M0 可靠闭环", "第 1–2 周", "A0 收尾、审计、快照、安全与回归", "参数/命令/指标一致，负向用例阻断"),
        ("M1 扩展底座", "第 2–3 周", "A1/A2/A3；v8 回归、v11 最小接入", "Linux 冒烟、目录边界、适配契约通过"),
        ("M2 在线内部版", "第 4–7 周", "B：认证、隔离、队列、GPU、SSE、历史", "并发/越权/重启恢复演练通过"),
        ("M3 多任务增量", "第 6–9 周", "C1 按业务优先级逐项交付", "每个任务真实数据集验收"),
        ("M4 结构实验一期", "第 8–12 周", "C2：解析、注册、变体、Agent、消融 UI", "可解释、可复现、可回退、有消融证据"),
        ("M5 CVAT 闭环", "第 10–12 周", "C3：快照导出、状态/产物、模型回传", "契约测试、权限与版本追踪通过"),
    ], [1450, 1900, 3460, 2550], 8.1)
    add_callout(doc, "排期说明", "M3、M4、M5 可在 B 的稳定接口完成后部分并行，因此时间点不是简单相加。若只有 1 人，按阶段串行执行；若 3 人，可让平台、ML 和 UI/集成并行，但不能跳过 Go/No-Go 门槛。", PALE_BLUE)

    page_break(doc)
    add_heading(doc, "13. 测试、质量与验收策略", 1)
    add_heading(doc, "13.1 分层测试", 2)
    add_table(doc, ["层级", "重点", "最小要求"], [
        ("单元测试", "参数类型/范围、规则诊断、状态机、适配映射", "关键分支和负向用例"),
        ("契约测试", "Repository、TaskAdapter、ModelAdapter、CVAT API", "接口升级不破坏领域层"),
        ("集成测试", "Worker+DB+GPU 租约、SSE 游标、取消/重试", "并发和恢复场景"),
        ("安全测试", "越权、路径穿越、注入、上传/目录边界", "高风险用例全部阻断"),
        ("真实训练冒烟", "最小真实数据、epochs=1、结果解析与审计", "每个任务/版本一条基线"),
        ("消融/质量评估", "baseline vs 变体、重复实验、资源成本", "C2 收益必须有证据"),
    ], [1600, 4160, 3600], 8.7)
    add_heading(doc, "13.2 总体验收指标", 2)
    for t in [
        "参数审计一致率 100%；非法 LLM 参数进入训练命令数量为 0。",
        "普通训练与调优任务均可追溯；产物可定位到数据快照、参数、代码和依赖版本。",
        "多用户越权测试全部阻止；GPU 分配冲突数量为 0。",
        "服务重启后任务状态可对账，历史事件可重放。",
        "结构实验必须拥有 baseline 和消融证据。",
    ]:
        add_bullet(doc, t)

    page_break(doc)
    add_heading(doc, "14. 建议的审批结论", 1)
    add_callout(doc, "建议批准：分阶段实施", "建议 Leader 批准以 2 名核心工程师为最低配置，先完成 M0/M1，再根据在线内部版的稳定性决定 C1/C2/C3 的并行程度。结构能力应立项，但不能提前绕过在线执行与实验审计底座。", GREEN)
    add_heading(doc, "14.1 本次需要批准的事项", 2)
    for t in [
        "确认一期边界：LLM 只修改白名单超参数；不在当前一期自动改模型结构。",
        "确认 A0 收尾与 A1/A2/A3 为下一批正式工作。",
        "确认在线内部可用版按 15–25 人日评估，而不是原先偏乐观的简单 Web 改造口径。",
        "确认结构解析与结构自调整后续同一期交付，预算 12–20 人日，并启用人工审批和消融验收。",
        "确认 CVAT 采用 API 解耦路线，首期优先完成单向数据快照交接。",
    ]:
        add_bullet(doc, t)
    add_heading(doc, "14.2 下一步（批准后）", 2)
    add_text(doc, "按当前协作方式分批实施：每批先更新测试，再实现功能，运行相关自动化检查和短 epoch 真实训练；每批结束向项目负责人汇报改动、测试证据、遗留风险和下一批范围。")
    add_text(doc, "文档依据", bold=True, color=DARK, size=10.5, after=4)
    add_bullet(doc, "docs/方案审查_20260814.md")
    add_bullet(doc, "docs/后续研发计划_20260814.md")
    add_bullet(doc, "docs/implementation_plan_20260814.md")
    add_bullet(doc, "docs/roadmap_20260814.md")

    # Prevent rows from splitting across pages and set document properties.
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
    doc.core_properties.title = "Auto-Tune 后续研发路线与实施评估"
    doc.core_properties.subject = "技术难度、资源、工期与分阶段验收"
    doc.core_properties.author = "Auto-Tune 项目组"
    doc.core_properties.keywords = "Auto-Tune, YOLO, LLM, 超参数调优, 模型结构, CVAT"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
